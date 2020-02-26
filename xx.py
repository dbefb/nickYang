from docx import Document
from docx.shared import Inches

document = Document()                          # 打开一个基于默认“模板”的空白文档
document.add_heading('Document Title', 0)      # 添加标题

p = document.add_paragraph('A plain paragraph having some ')    # 增加一段
p.add_run('bold').bold = True                                   # 设置样式
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('./gggg/001.png', width=Inches(1.25))     # 插入图片（默认居左）

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3) # 添加表
hdr_cells = table.rows[0].cells            # 设置表首行标题
hdr_cells[0].text = 'Qty'                  # 表首行标题赋值
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:              # for循环将records内容赋值到单元格内
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()                 # 添加分页符

document.save('./demo.docx')              # 保存文件
