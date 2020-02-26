import random
from docx import Document
from docx.shared import Inches

class gradeOne(object):
    def __init__(self,grade,cycleIndex):
        self.operator=['+','-']
        self.cycleIndex=cycleIndex
        self.grade=grade
    def formula(self):
        formula=str(random.randint(1,10))+self.operator[random.randint(0,1)]+str(random.randint(1,10))+self.operator[random.randint(0,1)]+str(random.randint(1,10))
        answer=eval(formula)

        return answer,formula
            
    def creatDocx(self):
        a=0
        switch=0
        switch2=0
        if self.grade=='三':
            switch3=3
        elif  self.grade=='五':
            switch3=2
        else:
            switch3=3
        document = Document()
        document2=Document()
        head=document.add_heading('小学{}年级四则运算'.format(self.grade), 0)      # 添加标题
        head.alignment = 1
        head=document2.add_heading('小学{}年级四则运算答案'.format(self.grade), 0)
        head.alignment = 1
        while a!=self.cycleIndex:
            try:
                answer,formula=self.formula()
            
                if answer >= 0 and float(answer).is_integer() == True:
                    # print(formula+'='+str(int(answer)))
                    if switch == 0:
                        p=document.add_paragraph('{}: {}='.format(a+1,formula))    # 增加一段
                        p2=document2.add_paragraph('{}: {}'.format(a+1,int(answer)))
                        p.alignment = 1
                        p2.alignment = 1
                        a+=1
                        switch=1
                        continue
                    else:
                        p.add_run('                  ')
                        p.add_run('{}: {}='.format(a+1,formula))
                        p2.add_run('                  ')
                        p2.add_run('{}: {}'.format(a+1,int(answer)))
                        
                        switch2+=1
                        if switch2 == switch3:
                            switch=0
                            switch2=0
                        a+=1
            except ZeroDivisionError:
                a-=1
        document.save('./小学{}年级四则运算.docx'.format(self.grade))              # 保存文件
        document2.save('./小学{}年级四则运算答案.docx'.format(self.grade))  

class gradeTwo(gradeOne):
    def __init__(self,grade,cycleIndex):
        self.operator=['+','-','*','/']
        self.operator2=['+','-']
        self.grade=grade
        self.cycleIndex=cycleIndex
    def formula(self):
        formula=str(random.randint(1,10))+self.operator[random.randint(0,3)]+str(random.randint(1,10))+self.operator2[random.randint(0,1)]+str(random.randint(1,10))
        answer=eval(formula)

        return answer,formula
            
class gradeThree(gradeOne):
    def __init__(self,grade,cycleIndex):
        self.operator=['+','-','*','/']
        self.operator2=['(','N']  #判断是否要加括号
        self.grade=grade
        self.cycleIndex=cycleIndex
        
    def formula(self):
        
        op_1=self.operator[random.randint(0,3)]
        op_2=self.operator[random.randint(0,3)]
        op_3=self.operator2[random.randint(0,1)]
        if (op_1 == '+' or op_1 == '-') and  op_3 == '(' and (op_2 == '*' or op_2 == '/'):  #判断是否要加括号
            formula='('+str(random.randint(1,100))+op_1+str(random.randint(1,100))+')'+op_2+str(random.randint(1,100))
        elif (op_1 == '*' or op_1 == '/') and  op_3 == '(' and (op_2 == '+' or op_2 == '-'):
            formula=str(random.randint(1,100))+op_1+'('+str(random.randint(1,100))+op_2+str(random.randint(1,100))+')'
        else :
            formula=str(random.randint(1,100))+op_1+str(random.randint(1,100))+op_2+str(random.randint(1,100))
        answer=eval(formula)

        return answer,formula
            
class gradeFour(gradeThree):
    
    def creatDocx(self):
        a=0
        switch=0
        switch2=0
        document = Document()
        document2=Document()
        head=document.add_heading('小学{}年级四则运算'.format(self.grade), 0)      # 添加标题
        head.alignment = 1
        head=document2.add_heading('小学{}年级四则运算答案'.format(self.grade), 0)
        head.alignment = 1
        while a!=self.cycleIndex:
            try:
                answer,formula=self.formula()
                formula=formula+self.operator[random.randint(0,3)]+str(random.randint(1,100))
                answer=eval(formula)
                if answer >= 0 and float(answer).is_integer() == True:
                    # print(formula+'='+str(int(answer)))
                    if switch == 0:
                        p=document.add_paragraph('{}: {}='.format(a+1,formula))    # 增加一段
                        p2=document2.add_paragraph('{}: {}'.format(a+1,int(answer)))
                        p.alignment = 1
                        p2.alignment = 1
                        a+=1
                        switch=1
                        continue
                    else:
                        p.add_run('                  ')
                        p.add_run('{}: {}='.format(a+1,formula))
                        p2.add_run('                  ')
                        p2.add_run('{}: {}'.format(a+1,int(answer)))
                        
                        switch2+=1
                        if switch2 == 2:
                            switch=0
                            switch2=0
                        a+=1
            except ZeroDivisionError:
                a-=1
        document.save('./小学{}年级四则运算.docx'.format(self.grade))                 # 保存文件
        document2.save('./小学{}年级四则运算答案.docx'.format(self.grade))  

class gradeFive(gradeThree):

    def formula(self):
        
        op_1=self.operator[random.randint(0,3)]
        op_2=self.operator[random.randint(0,3)]
        op_3=self.operator2[random.randint(0,1)]
        if (op_1 == '+' or op_1 == '-') and  op_3 == '(' and (op_2 == '*' or op_2 == '/'):  #判断是否要加括号
            formula='('+format(random.uniform(1,100),'.1f')+op_1+format(random.uniform(1,10),'.1f')+')'+op_2+format(random.uniform(1,100),'.1f')
        elif (op_1 == '*' or op_1 == '/') and  op_3 == '(' and (op_2 == '+' or op_2 == '-'):
            formula=format(random.uniform(1,100),'.1f')+op_1+'('+format(random.uniform(1,10),'.1f')+op_2+format(random.uniform(1,100),'.1f')+')'
        else :
            formula=format(random.uniform(1,100),'.1f')+op_1+format(random.uniform(1,10),'.1f')+op_2+format(random.uniform(1,100),'.1f')
        answer=eval(formula)
        
        return answer,formula

class gradeSix(gradeFive):

    def formula(self):
        
        op_1=self.operator[random.randint(0,3)]
        op_2=self.operator[random.randint(0,3)]
        op_3=self.operator2[random.randint(0,1)]
        if (op_1 == '+' or op_1 == '-') and  op_3 == '(' and (op_2 == '*' or op_2 == '/'):  #判断是否要加括号
            formula='('+format(random.uniform(1,100),'.1f')+op_1+format(random.uniform(1,100),'.1f')+')'+op_2+format(random.uniform(1,100),'.1f')
        elif (op_1 == '*' or op_1 == '/') and  op_3 == '(' and (op_2 == '+' or op_2 == '-'):
            formula=format(random.uniform(1,100),'.1f')+op_1+'('+format(random.uniform(1,100),'.1f')+op_2+format(random.uniform(1,100),'.1f')+')'
        else :
            formula=format(random.uniform(1,100),'.1f')+op_1+format(random.uniform(1,100),'.1f')+op_2+format(random.uniform(1,100),'.1f')
        answer=eval(formula)
        
        return answer,formula

    def creatDocx(self):
        a=0
        switch=0
        switch2=0
        document = Document()
        document2=Document()
        head=document.add_heading('小学{}年级四则运算'.format(self.grade), 0)      # 添加标题
        head.alignment = 1
        head=document2.add_heading('小学{}年级四则运算答案'.format(self.grade), 0)
        head.alignment = 1
        while a!=self.cycleIndex:
            try:
                answer,formula=self.formula()
                answer2,formula2=self.formula()
                formula=formula+self.operator[random.randint(0,3)]+formula2
                answer=eval(formula)
                if   float(answer).is_integer() == True:
                    # print(formula+'='+str(int(answer)))
                    if switch == 0:
                        p=document.add_paragraph('{}: {}='.format(a,formula))     # 增加一段
                        p2=document2.add_paragraph('{}: {}'.format(a+1,int(answer)))
                        p.alignment = 1
                        p2.alignment = 1
                        a+=1
                        switch=1
                        continue
                    else:
                        p.add_run('                  ')
                        p.add_run('{}: {}='.format(a+1,formula))
                        p2.add_run('                  ')
                        p2.add_run('{}: {}'.format(a+1,int(answer)))

                        switch2+=1
                        if switch2 == 1:
                            switch=0
                            switch2=0
                        a+=1
            except ZeroDivisionError:
                a-=1
        document.save('./小学{}年级四则运算.docx'.format(self.grade))               # 保存文件
        document2.save('./小学{}年级四则运算答案.docx'.format(self.grade))  

if __name__ == '__main__':
    while 1:
        grade=input('请输入年级:')
        cycleIndex=input('请输入题目数量:')
        if int(grade) == 1:
            grade_One=gradeOne('一',int(cycleIndex))
            grade_One.creatDocx()

        elif int(grade) == 2:
            grade_Two=gradeTwo('二',int(cycleIndex))
            grade_Two.creatDocx()

        elif int(grade) == 3:
            grade_Three=gradeThree('三',int(cycleIndex))
            grade_Three.creatDocx()

        elif int(grade) == 4:
            grade_Four=gradeFour('四',int(cycleIndex))
            grade_Four.creatDocx()

        elif int(grade) == 5:
            grade_Five=gradeFive('五',int(cycleIndex))
            grade_Five.creatDocx()
            
        elif int(grade) == 6:
            grade_Six=gradeSix('六',int(cycleIndex))
            grade_Six.creatDocx()